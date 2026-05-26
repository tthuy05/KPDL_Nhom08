from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUT_FILE = Path("BaoCao_DoAn_Nhom08.docx")


def set_run_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text="", align=None, bold=False, italic=False, size=13):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_run_font(r, size=15 if level == 1 else 14, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        r = hdr[i].paragraphs[0].add_run(h)
        set_run_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            r = cells[i].paragraphs[0].add_run(str(value))
            set_run_font(r)
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(6)


def build_report():
    doc = Document()
    configure_document(doc)

    # Trang bìa
    add_paragraph(doc, "TRƯỜNG: [Tên trường]", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "KHOA: [Tên khoa]", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "MÔN HỌC: KHAI PHÁ DỮ LIỆU", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc)
    add_paragraph(doc, "BÁO CÁO ĐỒ ÁN MÔN HỌC", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    add_paragraph(
        doc,
        "ỨNG DỤNG KỸ THUẬT GOM NHÓM TRONG KHAI THÁC DỮ LIỆU KHÁCH HÀNG ĐỂ PHÂN KHÚC THỊ TRƯỜNG",
        WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=16,
    )
    add_paragraph(doc)
    add_paragraph(doc, "Nhóm thực hiện: Nhóm 08", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Giảng viên hướng dẫn: [Tên giảng viên]", WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Lớp: [Tên lớp]", WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Năm học: 2025 - 2026", WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc)
    add_table(
        doc,
        ["STT", "Họ và tên", "MSSV", "Nhiệm vụ"],
        [
            ["1", "[Tên sinh viên 1]", "[MSSV]", "Thu thập dữ liệu, tiền xử lý, báo cáo"],
            ["2", "[Tên sinh viên 2]", "[MSSV]", "Xây dựng mô hình, giao diện, kiểm thử"],
            ["3", "[Tên sinh viên 3]", "[MSSV]", "Trực quan hóa, slide, thuyết trình"],
        ],
    )
    doc.add_page_break()

    add_heading(doc, "MỤC LỤC", 1)
    for item in [
        "1. Tổng quan đề tài",
        "2. Cơ sở lý thuyết",
        "3. Thu thập và mô tả dữ liệu",
        "4. Tiền xử lý dữ liệu",
        "5. Mô hình KMeans và đánh giá",
        "6. Xây dựng ứng dụng",
        "7. Kết quả thực nghiệm và phân tích khách hàng",
        "8. Kết luận và hướng phát triển",
        "9. Phân công công việc nhóm",
        "10. Phụ lục hướng dẫn chạy chương trình",
    ]:
        add_paragraph(doc, item)
    doc.add_page_break()

    add_heading(doc, "1. TỔNG QUAN ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Lý do chọn đề tài", 2)
    add_paragraph(
        doc,
        "Trong hoạt động kinh doanh, khách hàng không có hành vi giống nhau. Một số khách hàng có thu nhập cao "
        "và chi tiêu nhiều, một số mua hàng thường xuyên nhưng giá trị chi tiêu thấp, trong khi nhóm khác ít tương tác "
        "với doanh nghiệp. Nếu doanh nghiệp áp dụng cùng một chính sách cho toàn bộ khách hàng thì hiệu quả marketing "
        "và chăm sóc khách hàng sẽ không cao.",
    )
    add_paragraph(
        doc,
        "Khai phá dữ liệu giúp phát hiện các mẫu ẩn trong dữ liệu khách hàng. Trong đó, kỹ thuật gom nhóm là hướng "
        "tiếp cận phù hợp khi dữ liệu chưa có nhãn phân loại sẵn. Đề tài này xây dựng hệ thống phân cụm khách hàng "
        "bằng thuật toán KMeans nhằm hỗ trợ phân khúc thị trường.",
    )
    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    for text in [
        "Thu thập và chuẩn bị bộ dữ liệu khách hàng phục vụ bài toán phân khúc thị trường.",
        "Thực hiện làm sạch dữ liệu, xử lý missing values, encode dữ liệu và chuẩn hóa đặc trưng.",
        "Áp dụng thuật toán KMeans để gom nhóm khách hàng có đặc điểm tương đồng.",
        "Sử dụng Elbow Method và các chỉ số đánh giá để lựa chọn, kiểm tra chất lượng phân cụm.",
        "Xây dựng ứng dụng Streamlit cho phép upload dữ liệu, chạy phân cụm, trực quan hóa và tải kết quả.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2. CƠ SỞ LÝ THUYẾT", 1)
    add_heading(doc, "2.1. Khai phá dữ liệu khách hàng", 2)
    add_paragraph(
        doc,
        "Khai phá dữ liệu là quá trình trích xuất tri thức, quy luật hoặc mẫu có ý nghĩa từ tập dữ liệu. "
        "Trong lĩnh vực khách hàng, khai phá dữ liệu hỗ trợ phân khúc thị trường, dự đoán rời bỏ, gợi ý sản phẩm "
        "và tối ưu chiến dịch marketing.",
    )
    add_heading(doc, "2.2. Gom nhóm dữ liệu", 2)
    add_paragraph(
        doc,
        "Gom nhóm là kỹ thuật học không giám sát. Các đối tượng được chia thành các nhóm sao cho đối tượng trong "
        "cùng nhóm có mức độ tương đồng cao, còn đối tượng thuộc các nhóm khác nhau có mức độ khác biệt lớn.",
    )
    add_heading(doc, "2.3. Thuật toán KMeans", 2)
    add_paragraph(
        doc,
        "KMeans là thuật toán phân cụm phổ biến, hoạt động bằng cách chia dữ liệu thành K cụm. Mỗi cụm được đại diện "
        "bởi một tâm cụm. Thuật toán lặp lại quá trình gán điểm dữ liệu vào tâm cụm gần nhất và cập nhật tâm cụm "
        "cho đến khi hội tụ.",
    )
    for text in [
        "Chọn số cụm K.",
        "Khởi tạo K tâm cụm ban đầu.",
        "Gán mỗi khách hàng vào cụm có tâm gần nhất theo khoảng cách Euclidean.",
        "Cập nhật tâm cụm bằng trung bình các điểm trong cụm.",
        "Lặp lại đến khi tâm cụm ổn định hoặc đạt số vòng lặp tối đa.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "2.4. Chuẩn hóa dữ liệu", 2)
    add_paragraph(
        doc,
        "KMeans dựa trên khoảng cách giữa các điểm dữ liệu. Nếu các cột có thang đo khác nhau, ví dụ thu nhập có "
        "giá trị lớn hơn nhiều so với điểm chi tiêu, cột thu nhập sẽ chi phối kết quả. Vì vậy dữ liệu được chuẩn hóa "
        "bằng StandardScaler trước khi đưa vào mô hình.",
    )
    add_heading(doc, "2.5. Chỉ số đánh giá", 2)
    add_table(
        doc,
        ["Chỉ số", "Ý nghĩa"],
        [
            ["Inertia", "Tổng khoảng cách bình phương từ các điểm đến tâm cụm. Giá trị càng nhỏ càng tốt."],
            ["Silhouette Score", "Đo mức độ tách biệt giữa các cụm. Giá trị càng gần 1 càng tốt."],
            ["Davies-Bouldin Index", "Đo mức độ chồng lấn giữa các cụm. Giá trị càng nhỏ càng tốt."],
            ["Calinski-Harabasz Score", "Đo độ tách biệt giữa cụm so với độ chặt trong cụm. Giá trị càng lớn càng tốt."],
        ],
    )

    add_heading(doc, "3. THU THẬP VÀ MÔ TẢ DỮ LIỆU", 1)
    add_heading(doc, "3.1. Nguồn dữ liệu", 2)
    add_paragraph(
        doc,
        "Dữ liệu chính dùng để kiểm thử hệ thống là bộ Customer Churn & Segmentation Dataset (Synthetic) trên Kaggle. "
        "Bộ dữ liệu mô phỏng hành vi khách hàng trong thương mại điện tử/bán lẻ, gồm 2.000 bản ghi và các thuộc tính "
        "liên quan đến nhân khẩu học, thu nhập, điểm chi tiêu, tần suất mua hàng và trạng thái rời bỏ.",
    )
    add_paragraph(
        doc,
        "Ngoài việc dùng dataset Kaggle mặc định, ứng dụng còn cho phép người dùng upload file CSV khác trực tiếp trên giao diện.",
    )
    add_heading(doc, "3.2. Thuộc tính sử dụng", 2)
    add_table(
        doc,
        ["Cột sau chuẩn hóa", "Cột gốc Kaggle", "Mô tả"],
        [
            ["CustomerID", "CustomerID", "Mã định danh khách hàng"],
            ["Gender", "Gender", "Giới tính khách hàng"],
            ["Age", "Age", "Tuổi khách hàng"],
            ["Income", "Annual_Income", "Thu nhập của khách hàng"],
            ["SpendingScore", "Spending_Score", "Điểm chi tiêu từ 1 đến 100"],
            ["PurchaseFrequency", "Purchase_Frequency", "Tần suất mua hàng"],
        ],
    )
    add_heading(doc, "3.3. Thống kê dữ liệu Kaggle", 2)
    add_table(
        doc,
        ["Thông tin", "Giá trị"],
        [
            ["Số dòng", "2.000"],
            ["Số cột dùng trong project", "6"],
            ["Missing values", "0"],
            ["Dòng trùng lặp", "0"],
            ["Nam", "1.012"],
            ["Nữ", "988"],
        ],
    )
    add_table(
        doc,
        ["Thuộc tính", "Min", "Mean", "Median", "Max"],
        [
            ["Age", "18", "43,57", "43", "69"],
            ["Income", "9.857", "59.849,51", "60.036,50", "110.053"],
            ["SpendingScore", "1", "50,47", "51", "100"],
            ["PurchaseFrequency", "1,00", "8,00", "8,00", "15,00"],
        ],
    )

    add_heading(doc, "4. TIỀN XỬ LÝ DỮ LIỆU", 1)
    add_paragraph(doc, "Pipeline tiền xử lý được xây dựng trong file preprocessing.py, gồm các bước:")
    for text in [
        "Kiểm tra các cột bắt buộc: Age, Income, SpendingScore, PurchaseFrequency.",
        "Chuyển các cột số về kiểu numeric để tránh lỗi dữ liệu dạng chuỗi.",
        "Xử lý missing values: cột số điền median, cột phân loại điền mode hoặc Unknown nếu toàn bộ cột bị thiếu.",
        "Xóa các dòng trùng lặp.",
        "Encode cột Gender thành Gender_Encoded bằng LabelEncoder.",
        "Chuẩn hóa các đặc trưng số bằng StandardScaler.",
    ]:
        add_bullet(doc, text)
    add_paragraph(
        doc,
        "Hệ thống có cơ chế kiểm tra lỗi khi người dùng upload CSV. Nếu thiếu cột bắt buộc, cột số toàn missing hoặc "
        "dữ liệu còn quá ít để chạy KMeans, ứng dụng sẽ hiển thị thông báo lỗi rõ ràng.",
    )
    add_paragraph(
        doc,
        "Các đặc trưng đưa vào mô hình gồm Age, Income, SpendingScore và PurchaseFrequency. Cột Gender_Encoded được tạo "
        "để minh họa bước encode dữ liệu phân loại và phục vụ mở rộng sau này.",
    )

    add_heading(doc, "5. MÔ HÌNH KMEANS VÀ ĐÁNH GIÁ", 1)
    add_heading(doc, "5.1. Lý do chọn KMeans", 2)
    add_paragraph(
        doc,
        "KMeans được chọn vì thuật toán đơn giản, dễ triển khai, tốc độ nhanh và phù hợp với bài toán phân cụm khách hàng "
        "khi dữ liệu đã được biểu diễn bằng các đặc trưng số.",
    )
    add_heading(doc, "5.2. Xác định số cụm K", 2)
    add_paragraph(
        doc,
        "Hệ thống chạy KMeans với nhiều giá trị K từ 2 đến 10, sau đó so sánh Inertia và Silhouette Score. "
        "Với dữ liệu Kaggle, Silhouette Score cao nhất trong khoảng thử nghiệm đạt được tại K = 8.",
    )
    add_table(
        doc,
        ["K", "Inertia", "Silhouette Score"],
        [
            ["2", "6483,26", "0,1870"],
            ["3", "5531,67", "0,1801"],
            ["4", "4802,66", "0,1943"],
            ["5", "4272,91", "0,1969"],
            ["6", "3878,53", "0,1985"],
            ["7", "3526,03", "0,2055"],
            ["8", "3193,34", "0,2145"],
            ["9", "2992,89", "0,2091"],
            ["10", "2798,43", "0,2084"],
        ],
    )
    add_heading(doc, "5.3. Kết quả đánh giá", 2)
    add_table(
        doc,
        ["Chỉ số", "Giá trị với K = 8"],
        [
            ["Silhouette Score", "0,2145"],
            ["Davies-Bouldin Index", "1,1664"],
            ["Calinski-Harabasz Score", "428,36"],
            ["Inertia", "3193,34"],
        ],
    )
    add_paragraph(
        doc,
        "Silhouette Score không quá cao, cho thấy dữ liệu khách hàng có mức độ chồng lấn giữa các nhóm. Tuy nhiên, "
        "kết quả vẫn có thể dùng để phân tích xu hướng và hỗ trợ phân khúc thị trường ở mức cơ bản.",
    )

    add_heading(doc, "6. XÂY DỰNG ỨNG DỤNG", 1)
    add_heading(doc, "6.1. Công nghệ sử dụng", 2)
    add_table(
        doc,
        ["Công nghệ/thư viện", "Vai trò"],
        [
            ["Python", "Ngôn ngữ lập trình chính"],
            ["Pandas", "Đọc, xử lý và biến đổi dữ liệu bảng"],
            ["NumPy", "Tính toán số học"],
            ["Scikit-learn", "StandardScaler, LabelEncoder, KMeans và chỉ số đánh giá"],
            ["Matplotlib", "Vẽ histogram, elbow chart và scatter plot"],
            ["Seaborn", "Vẽ heatmap tương quan"],
            ["Streamlit", "Xây dựng dashboard tương tác"],
        ],
    )
    add_heading(doc, "6.2. Cấu trúc project", 2)
    add_table(
        doc,
        ["File/thư mục", "Chức năng"],
        [
            ["data/", "Lưu dataset Kaggle gốc và dataset Kaggle đã chuẩn hóa"],
            ["preprocessing.py", "Tiền xử lý dữ liệu"],
            ["clustering.py", "Elbow Method, KMeans, đánh giá và phân tích cụm"],
            ["visualization.py", "Các hàm trực quan hóa dữ liệu"],
            ["app.py", "Ứng dụng Streamlit cho người dùng cuối"],
            ["requirements.txt", "Danh sách thư viện cần cài đặt"],
        ],
    )
    add_heading(doc, "6.3. Chức năng giao diện", 2)
    for text in [
        "Upload CSV hoặc dùng dataset Kaggle mặc định.",
        "Hiển thị tổng quan dữ liệu: số khách hàng, số missing values và số duplicate.",
        "Vẽ Histogram để quan sát phân phối từng đặc trưng.",
        "Vẽ Heatmap để quan sát tương quan giữa các biến.",
        "Chọn số cụm K bằng slider.",
        "Vẽ Elbow Method và Silhouette Score để hỗ trợ chọn K.",
        "Chạy KMeans và hiển thị các chỉ số đánh giá.",
        "Vẽ Scatter Plot theo Income - SpendingScore và Age - PurchaseFrequency.",
        "Gán tên mô tả cho từng cụm khách hàng.",
        "Hiển thị bảng kết quả và tải file CSV sau phân cụm.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "6.4. Luồng xử lý hệ thống", 2)
    add_paragraph(
        doc,
        "Dữ liệu đầu vào -> Kiểm tra cột bắt buộc -> Xử lý missing values -> Xóa duplicate -> Encode Gender -> "
        "Feature scaling -> Xác định K -> Huấn luyện KMeans -> Gán nhãn cluster -> Phân tích nhóm khách hàng -> "
        "Đánh giá mô hình -> Hiển thị và tải kết quả.",
    )
    add_paragraph(
        doc,
        "Gợi ý: chèn thêm ảnh giao diện dashboard, Histogram, Heatmap, Elbow/Silhouette và Scatter Plot vào mục này.",
        italic=True,
    )

    add_heading(doc, "7. KẾT QUẢ THỰC NGHIỆM VÀ PHÂN TÍCH KHÁCH HÀNG", 1)
    add_heading(doc, "7.1. Kết quả phân cụm", 2)
    add_paragraph(
        doc,
        "Với dữ liệu Kaggle đã chuẩn hóa, hệ thống chọn K = 8 theo Silhouette Score trong khoảng K từ 2 đến 10. "
        "Số lượng khách hàng trong từng cụm khá cân bằng.",
    )
    add_table(
        doc,
        ["Cụm", "Số khách hàng", "Tỷ lệ"],
        [
            ["0", "232", "11,6%"],
            ["1", "251", "12,6%"],
            ["2", "266", "13,3%"],
            ["3", "237", "11,8%"],
            ["4", "240", "12,0%"],
            ["5", "243", "12,2%"],
            ["6", "267", "13,4%"],
            ["7", "264", "13,2%"],
        ],
    )
    add_heading(doc, "7.2. Diễn giải nhóm khách hàng", 2)
    add_table(
        doc,
        ["Nhóm khách hàng", "Số lượng", "Đặc điểm tổng quát", "Gợi ý chiến lược"],
        [
            ["Khách tiềm năng", "747", "Thu nhập và/hoặc điểm chi tiêu ở mức khá.", "Ưu đãi cá nhân hóa, chương trình thành viên."],
            ["Khách chi tiêu thấp", "770", "Điểm chi tiêu hoặc mức đóng góp chưa cao.", "Combo sản phẩm, gợi ý sản phẩm phù hợp ngân sách."],
            ["Khách ít mua hàng", "483", "Tần suất mua thấp, ít tương tác.", "Voucher quay lại, nhắc nhở, khảo sát lý do ít mua."],
        ],
    )
    add_paragraph(
        doc,
        "KMeans chỉ gán nhãn cụm dạng số. Tên nhóm khách hàng là bước diễn giải nghiệp vụ dựa trên thu nhập, điểm chi tiêu "
        "và tần suất mua hàng trung bình của từng cụm.",
    )
    add_heading(doc, "7.3. Ý nghĩa biểu đồ trong ứng dụng", 2)
    add_table(
        doc,
        ["Biểu đồ", "Ý nghĩa", "Dựa vào đâu"],
        [
            ["Histogram", "Cho biết phân phối tuổi, thu nhập, điểm chi tiêu và tần suất mua.", "Dựa trên từng cột dữ liệu sau tiền xử lý."],
            ["Heatmap", "Cho biết mức độ tương quan giữa các biến số.", "Dựa trên hệ số tương quan Pearson."],
            ["Elbow Method", "Quan sát điểm mà Inertia bắt đầu giảm chậm khi tăng K.", "Dựa trên khoảng cách từ điểm dữ liệu đến tâm cụm."],
            ["Silhouette Chart", "So sánh chất lượng phân cụm theo từng K.", "Dựa trên độ gần của điểm với cụm của nó so với cụm khác."],
            ["Scatter Plot", "Quan sát các cụm khách hàng trên hai chiều dễ nhìn.", "Dựa trên nhãn cụm KMeans."],
        ],
    )

    add_heading(doc, "8. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_paragraph(
        doc,
        "Đồ án đã xây dựng được quy trình khai phá dữ liệu khách hàng phục vụ phân khúc thị trường bằng kỹ thuật gom nhóm. "
        "Hệ thống thực hiện đầy đủ các bước từ chuẩn bị dữ liệu, làm sạch dữ liệu, xử lý missing values, encode dữ liệu, "
        "feature scaling, xác định số cụm, huấn luyện KMeans, đánh giá mô hình và trực quan hóa kết quả trên giao diện Streamlit.",
    )
    add_heading(doc, "8.1. Hạn chế", 2)
    for text in [
        "Thuật toán chính mới là KMeans, chưa so sánh sâu với DBSCAN hoặc Agglomerative Clustering.",
        "Quy tắc đặt tên cụm còn dựa trên ngưỡng thủ công.",
        "Dữ liệu Kaggle là dữ liệu synthetic, chưa phản ánh hoàn toàn dữ liệu doanh nghiệp thực tế.",
        "Chưa đưa các biến phân loại như Membership_Level vào mô hình.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "8.2. Hướng phát triển", 2)
    for text in [
        "Thêm so sánh với các thuật toán phân cụm khác.",
        "Bổ sung kiểm tra outlier và xử lý giá trị bất thường.",
        "Cho phép người dùng chọn feature dùng để phân cụm ngay trên giao diện.",
        "Bổ sung xuất báo cáo tự động sau phân cụm.",
        "Kết hợp phân cụm với bài toán dự đoán churn.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "9. PHÂN CÔNG CÔNG VIỆC NHÓM", 1)
    add_table(
        doc,
        ["Thành viên", "Nhiệm vụ", "Mức độ hoàn thành"],
        [
            ["[Tên sinh viên 1]", "Tìm hiểu đề tài, thu thập dữ liệu, mô tả dataset.", "100%"],
            ["[Tên sinh viên 2]", "Xây dựng pipeline tiền xử lý, KMeans, đánh giá mô hình.", "100%"],
            ["[Tên sinh viên 3]", "Xây dựng giao diện Streamlit, trực quan hóa, slide và demo.", "100%"],
        ],
    )
    add_paragraph(doc, "Cần thay placeholder bằng tên thật, MSSV và nhiệm vụ thực tế trước khi nộp.", italic=True)

    add_heading(doc, "10. PHỤ LỤC HƯỚNG DẪN CHẠY CHƯƠNG TRÌNH", 1)
    add_heading(doc, "10.1. Cài đặt thư viện", 2)
    add_paragraph(doc, "pip install -r requirements.txt")
    add_heading(doc, "10.2. Chạy ứng dụng Streamlit", 2)
    add_paragraph(doc, "streamlit run app.py")
    add_paragraph(doc, "Sau đó mở trình duyệt tại địa chỉ http://localhost:8501.")
    add_heading(doc, "10.3. Dữ liệu sử dụng khi demo", 2)
    for text in [
        "customer_data_with_churn.csv: dữ liệu Kaggle gốc.",
        "customers_kaggle.csv: dữ liệu Kaggle đã đổi tên cột để phù hợp với pipeline (dataset mặc định).",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "10.4. Các file cần nộp", 2)
    for text in [
        "File báo cáo Word.",
        "File PowerPoint thuyết trình.",
        "Toàn bộ source code.",
        "Dataset sử dụng.",
        "Nén thư mục theo tên Nhom08_DoAn.zip.",
    ]:
        add_bullet(doc, text)

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_report()
    print(OUT_FILE.resolve())
